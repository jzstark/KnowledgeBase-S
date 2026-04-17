"""
Word Source — 使用 python-docx 提取 .docx 文档正文。
"""

import docx

from sources.base import RawItem
from sources.file_base import FileSourceMixin


class WordSource(FileSourceMixin):
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(self, source_id: str, uploads: list[dict]):
        self.source_id = source_id
        self.uploads = uploads

    def extract_text(self, raw: RawItem) -> str:
        try:
            doc = docx.Document(raw.raw_ref["path"])

            # Word 核心属性标题优先，其次找第一个 Heading 样式段落
            meta_title = (doc.core_properties.title or "").strip()
            if meta_title:
                raw.title = meta_title
            else:
                for p in doc.paragraphs:
                    if p.text.strip() and "heading" in p.style.name.lower():
                        raw.title = p.text.strip()[:120]
                        break

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs).strip()
        except Exception:
            return ""
